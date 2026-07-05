#!/usr/bin/env python3
"""
Convertidor de Markdown a DOCX Profesional

Convierte INFORME_ALERTA_TEMPRANA_v2.md a un documento Word con formato
institucional de la Universidad Autónoma de Chile.

Uso:
    python scripts/convert_md_to_docx.py
    python scripts/convert_md_to_docx.py --ascii-as-image

Salida:
    data/report/INFORME_ALERTA_TEMPRANA_v2.docx
"""

import argparse
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Tuple, Union
from enum import Enum

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# =============================================================================
# CONFIGURACIÓN
# =============================================================================

BASE_DIR = Path(__file__).parent.parent
REPORT_DIR = BASE_DIR / "data" / "report"
INPUT_FILE = REPORT_DIR / "INFORME_ALERTA_TEMPRANA_v2.md"
OUTPUT_FILE = REPORT_DIR / "INFORME_ALERTA_TEMPRANA_v2.docx"
LOGO_FILE = BASE_DIR / "LOGO-UA-color.jpg"

# Colores corporativos
PRIMARY_COLOR = RGBColor(0x1E, 0x3A, 0x5F)      # Azul oscuro
SECONDARY_COLOR = RGBColor(0x2E, 0x86, 0xAB)    # Azul medio
ACCENT_COLOR = RGBColor(0xF1, 0x8F, 0x01)       # Naranja
TEXT_COLOR = RGBColor(0x30, 0x30, 0x30)         # Gris oscuro
WHITE_COLOR = RGBColor(0xFF, 0xFF, 0xFF)        # Blanco
GRAY_COLOR = RGBColor(0x60, 0x60, 0x60)         # Gris medio

# Hex para OxmlElement (sin #)
PRIMARY_HEX = '1E3A5F'
SECONDARY_HEX = '2E86AB'
ACCENT_HEX = 'F18F01'
LIGHT_GRAY_HEX = 'F5F5F5'
BORDER_GRAY_HEX = 'CCCCCC'
BLOCKQUOTE_BG_HEX = 'FFF8F0'

# Configuración de fuentes
FONT_NAME = 'Arial'
FONT_NAME_MONO = 'Consolas'

# =============================================================================
# ELEMENTOS DE MARKDOWN
# =============================================================================

class ElementType(Enum):
    HEADING = "heading"
    PARAGRAPH = "paragraph"
    TABLE = "table"
    CODE_BLOCK = "code_block"
    BLOCKQUOTE = "blockquote"
    IMAGE = "image"
    BULLET_LIST = "bullet_list"
    NUMBERED_LIST = "numbered_list"
    HORIZONTAL_RULE = "horizontal_rule"
    EMPTY = "empty"


@dataclass
class HeadingElement:
    level: int
    text: str
    type: ElementType = field(default=ElementType.HEADING)


@dataclass
class ParagraphElement:
    text: str
    type: ElementType = field(default=ElementType.PARAGRAPH)


@dataclass
class TableElement:
    headers: List[str]
    rows: List[List[str]]
    type: ElementType = field(default=ElementType.TABLE)


@dataclass
class CodeBlockElement:
    lines: List[str]
    language: str = ""
    type: ElementType = field(default=ElementType.CODE_BLOCK)


@dataclass
class BlockquoteElement:
    text: str
    type: ElementType = field(default=ElementType.BLOCKQUOTE)


@dataclass
class ImageElement:
    src: str
    alt: str
    type: ElementType = field(default=ElementType.IMAGE)


@dataclass
class ListElement:
    items: List[str]
    is_numbered: bool = False
    type: ElementType = field(default=ElementType.BULLET_LIST)


@dataclass
class HorizontalRuleElement:
    type: ElementType = field(default=ElementType.HORIZONTAL_RULE)


# Tipo union para todos los elementos
Element = Union[
    HeadingElement, ParagraphElement, TableElement, CodeBlockElement,
    BlockquoteElement, ImageElement, ListElement, HorizontalRuleElement
]


# =============================================================================
# FUNCIONES AUXILIARES PARA FORMATO WORD
# =============================================================================

def add_paragraph_border(paragraph, color: str, position: str = 'bottom', size: str = '6'):
    """Agrega borde a un párrafo."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    border = OxmlElement(f'w:{position}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), size)
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), color)
    pBdr.append(border)
    pPr.append(pBdr)


def add_left_border(paragraph, color: str, width: str = '24'):
    """Agrega borde izquierdo a un párrafo (para blockquotes)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), width)
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)


def set_paragraph_shading(paragraph, color: str):
    """Establece color de fondo en párrafo."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)


def set_table_borders(table, color: str):
    """Establece bordes coloridos en tabla."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tblBorders.append(border)

    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def set_cell_shading(cell, color_hex: str):
    """Establece color de fondo en celda."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=50, bottom=50, left=100, right=100):
    """Establece márgenes en celda (en twips)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def add_page_number_field(paragraph):
    """Agrega campo de número de página dinámico."""
    run = paragraph.add_run()
    run.font.name = FONT_NAME
    run.font.size = Pt(9)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# =============================================================================
# PARSER DE MARKDOWN
# =============================================================================

class MarkdownParser:
    """Parser de Markdown usando máquina de estados."""

    def __init__(self, content: str):
        self.content = content
        self.lines = content.split('\n')
        self.elements: List[Element] = []
        self.current_index = 0

    def parse(self) -> List[Element]:
        """Parsea el contenido completo."""
        while self.current_index < len(self.lines):
            line = self.lines[self.current_index]

            # Bloque de código
            if line.startswith('```'):
                self._parse_code_block()
            # Tabla
            elif line.startswith('|') and '|' in line[1:]:
                self._parse_table()
            # Encabezado
            elif re.match(r'^#{1,6}\s+', line):
                self._parse_heading()
            # Blockquote
            elif line.startswith('> '):
                self._parse_blockquote()
            # Imagen
            elif re.match(r'^!\[', line):
                self._parse_image()
            # Lista con viñetas
            elif re.match(r'^[\-\*]\s+', line):
                self._parse_bullet_list()
            # Lista numerada
            elif re.match(r'^\d+\.\s+', line):
                self._parse_numbered_list()
            # Línea horizontal
            elif line.strip() == '---':
                self.elements.append(HorizontalRuleElement(type=ElementType.HORIZONTAL_RULE))
                self.current_index += 1
            # Línea vacía
            elif not line.strip():
                self.current_index += 1
            # Párrafo
            else:
                self._parse_paragraph()

        return self.elements

    def _parse_heading(self):
        """Parsea encabezado."""
        line = self.lines[self.current_index]
        match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if match:
            level = len(match.group(1))
            text = match.group(2).strip()
            self.elements.append(HeadingElement(level=level, text=text))
        self.current_index += 1

    def _parse_paragraph(self):
        """Parsea párrafo (puede ser multilínea)."""
        lines = []
        while self.current_index < len(self.lines):
            line = self.lines[self.current_index]
            # Terminar párrafo si encontramos elemento especial
            if (not line.strip() or
                line.startswith('#') or
                line.startswith('```') or
                line.startswith('|') or
                line.startswith('> ') or
                line.startswith('- ') or
                line.startswith('* ') or
                re.match(r'^\d+\.\s+', line) or
                re.match(r'^!\[', line) or
                line.strip() == '---'):
                break
            lines.append(line)
            self.current_index += 1

        if lines:
            text = ' '.join(lines)
            self.elements.append(ParagraphElement(text=text))

    def _parse_table(self):
        """Parsea tabla Markdown."""
        table_lines = []
        while self.current_index < len(self.lines):
            line = self.lines[self.current_index]
            if line.startswith('|'):
                table_lines.append(line)
                self.current_index += 1
            else:
                break

        if len(table_lines) >= 2:
            # Primera línea = headers
            headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]

            # Segunda línea = separador (ignorar)
            # Resto = filas de datos
            rows = []
            for line in table_lines[2:]:
                cells = [c.strip() for c in line.split('|')[1:-1]]
                if cells:
                    rows.append(cells)

            self.elements.append(TableElement(headers=headers, rows=rows))

    def _parse_code_block(self):
        """Parsea bloque de código."""
        first_line = self.lines[self.current_index]
        language = first_line[3:].strip()  # Después de ```
        self.current_index += 1

        code_lines = []
        while self.current_index < len(self.lines):
            line = self.lines[self.current_index]
            if line.startswith('```'):
                self.current_index += 1
                break
            code_lines.append(line)
            self.current_index += 1

        self.elements.append(CodeBlockElement(lines=code_lines, language=language))

    def _parse_blockquote(self):
        """Parsea blockquote (puede ser multilínea)."""
        lines = []
        while self.current_index < len(self.lines):
            line = self.lines[self.current_index]
            if line.startswith('> '):
                lines.append(line[2:])  # Quitar "> "
                self.current_index += 1
            elif line.startswith('>'):
                lines.append(line[1:])  # Quitar ">"
                self.current_index += 1
            else:
                break

        text = ' '.join(lines)
        self.elements.append(BlockquoteElement(text=text))

    def _parse_image(self):
        """Parsea imagen ![alt](src)."""
        line = self.lines[self.current_index]
        match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', line)
        if match:
            alt = match.group(1)
            src = match.group(2)
            self.elements.append(ImageElement(src=src, alt=alt))
        self.current_index += 1

    def _parse_bullet_list(self):
        """Parsea lista con viñetas."""
        items = []
        while self.current_index < len(self.lines):
            line = self.lines[self.current_index]
            match = re.match(r'^[\-\*]\s+(.+)$', line)
            if match:
                items.append(match.group(1))
                self.current_index += 1
            else:
                break

        if items:
            elem = ListElement(items=items, is_numbered=False)
            elem.type = ElementType.BULLET_LIST
            self.elements.append(elem)

    def _parse_numbered_list(self):
        """Parsea lista numerada."""
        items = []
        while self.current_index < len(self.lines):
            line = self.lines[self.current_index]
            match = re.match(r'^\d+\.\s+(.+)$', line)
            if match:
                items.append(match.group(1))
                self.current_index += 1
            else:
                break

        if items:
            elem = ListElement(items=items, is_numbered=True)
            elem.type = ElementType.NUMBERED_LIST
            self.elements.append(elem)


# =============================================================================
# CONSTRUCTOR DE DOCUMENTO
# =============================================================================

class DocumentBuilder:
    """Construye el documento Word base."""

    def __init__(self, output_path: Path):
        self.output_path = output_path
        self.doc = Document()
        self._configure_page_layout()
        self._configure_styles()

    def _configure_page_layout(self):
        """Configura tamaño A4 y márgenes."""
        for section in self.doc.sections:
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2)

    def _configure_styles(self):
        """Configura estilos base del documento."""
        # Estilo Normal
        style = self.doc.styles['Normal']
        style.font.name = FONT_NAME
        style.font.size = Pt(11)
        style.font.color.rgb = TEXT_COLOR
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

    def add_title_page(self, title: str, subtitle: str, date: str, version: str):
        """Crea portada profesional."""
        # Espaciado superior
        for _ in range(3):
            self.doc.add_paragraph()

        # Logo
        if LOGO_FILE.exists():
            p_logo = self.doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_logo.add_run()
            run.add_picture(str(LOGO_FILE), width=Cm(8))

        # Espaciado
        for _ in range(3):
            self.doc.add_paragraph()

        # Línea decorativa
        p_line = self.doc.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_line.add_run('─' * 50)
        run.font.color.rgb = PRIMARY_COLOR

        # Título principal
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_title.add_run(title)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(26)
        run.font.color.rgb = PRIMARY_COLOR

        # Subtítulo
        p_subtitle = self.doc.add_paragraph()
        p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_subtitle.add_run(subtitle)
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.color.rgb = SECONDARY_COLOR

        # Espaciado
        for _ in range(2):
            self.doc.add_paragraph()

        # Línea decorativa inferior
        p_line2 = self.doc.add_paragraph()
        p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_line2.add_run('─' * 50)
        run.font.color.rgb = PRIMARY_COLOR

        # Información adicional
        for _ in range(4):
            self.doc.add_paragraph()

        p_info = self.doc.add_paragraph()
        p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_info.add_run(f"Fecha: {date}")
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR

        p_version = self.doc.add_paragraph()
        p_version.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_version.add_run(f"Versión: {version}")
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR

        # Salto de página
        self.doc.add_page_break()

    def add_table_of_contents(self, headings: List[HeadingElement]):
        """Genera tabla de contenidos manual."""
        # Título TOC
        p_toc = self.doc.add_paragraph()
        run = p_toc.add_run("Tabla de Contenidos")
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(18)
        run.font.color.rgb = PRIMARY_COLOR
        add_paragraph_border(p_toc, PRIMARY_HEX)

        self.doc.add_paragraph()

        # Entradas - usar el texto del heading directamente (ya incluye numeración)
        for heading in headings:
            # Saltar el heading "Tabla de Contenidos" del markdown original
            if 'tabla de contenidos' in heading.text.lower():
                continue

            if heading.level == 1:
                p = self.doc.add_paragraph()
                run = p.add_run(heading.text)
                run.bold = True
                run.font.name = FONT_NAME
                run.font.size = Pt(12)
                run.font.color.rgb = PRIMARY_COLOR
            elif heading.level == 2:
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.75)
                run = p.add_run(heading.text)
                run.font.name = FONT_NAME
                run.font.size = Pt(11)
                run.font.color.rgb = TEXT_COLOR

        self.doc.add_page_break()

    def setup_headers_footers(self, title: str):
        """Configura encabezados y pies de página."""
        for section in self.doc.sections:
            section.different_first_page_header_footer = True

            # Header (páginas 2+)
            header = section.header
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.clear()
            run = p.add_run(title[:60])
            run.font.name = FONT_NAME
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY_COLOR
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_paragraph_border(p, BORDER_GRAY_HEX)

            # Primera página sin header
            first_header = section.first_page_header
            if first_header.paragraphs:
                first_header.paragraphs[0].clear()

            # Footer
            footer = section.footer
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            run = p.add_run("Universidad Autónoma de Chile | Página ")
            run.font.name = FONT_NAME
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY_COLOR

            add_page_number_field(p)

    def save(self):
        """Guarda el documento."""
        self.doc.save(self.output_path)


# =============================================================================
# RENDERIZADOR DE ELEMENTOS
# =============================================================================

class ElementRenderer:
    """Renderiza elementos de Markdown a Word."""

    def __init__(self, doc: Document, base_image_path: Path, ascii_as_image: bool = False):
        self.doc = doc
        self.base_image_path = base_image_path
        self.ascii_as_image = ascii_as_image
        self.section_count = 0
        self.first_h1 = True

    def render(self, element: Element):
        """Renderiza un elemento según su tipo."""
        if element.type == ElementType.HEADING:
            self._render_heading(element)
        elif element.type == ElementType.PARAGRAPH:
            self._render_paragraph(element)
        elif element.type == ElementType.TABLE:
            self._render_table(element)
        elif element.type == ElementType.CODE_BLOCK:
            self._render_code_block(element)
        elif element.type == ElementType.BLOCKQUOTE:
            self._render_blockquote(element)
        elif element.type == ElementType.IMAGE:
            self._render_image(element)
        elif element.type in (ElementType.BULLET_LIST, ElementType.NUMBERED_LIST):
            self._render_list(element)
        elif element.type == ElementType.HORIZONTAL_RULE:
            self._render_horizontal_rule()

    def _add_formatted_text(self, paragraph, text: str, base_color: RGBColor = TEXT_COLOR):
        """Agrega texto con formato (negrita, cursiva)."""
        # Patrón para **bold** y *italic*
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*)'
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            if part.startswith('**') and part.endswith('**'):
                # Negrita
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Cursiva
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(part)

            run.font.name = FONT_NAME
            run.font.size = Pt(11)
            run.font.color.rgb = base_color

    def _render_heading(self, element: HeadingElement):
        """Renderiza encabezado."""
        # Salto de página antes de H1 (excepto el primero después de TOC)
        if element.level == 1:
            if not self.first_h1:
                self.doc.add_page_break()
            else:
                self.first_h1 = False
            self.section_count += 1

        # Configuración por nivel
        config = {
            1: {'size': Pt(18), 'color': PRIMARY_COLOR, 'border': True},
            2: {'size': Pt(14), 'color': SECONDARY_COLOR, 'border': False},
            3: {'size': Pt(12), 'color': PRIMARY_COLOR, 'border': False},
        }
        cfg = config.get(element.level, {'size': Pt(11), 'color': PRIMARY_COLOR, 'border': False})

        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12 if element.level > 1 else 0)
        p.paragraph_format.space_after = Pt(6)

        run = p.add_run(element.text)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = cfg['size']
        run.font.color.rgb = cfg['color']

        if cfg['border']:
            add_paragraph_border(p, PRIMARY_HEX)

    def _render_paragraph(self, element: ParagraphElement):
        """Renderiza párrafo."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._add_formatted_text(p, element.text)

    def _render_table(self, element: TableElement):
        """Renderiza tabla con estilo profesional."""
        if not element.headers:
            return

        num_cols = len(element.headers)
        table = self.doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        set_table_borders(table, SECONDARY_HEX)

        # Header
        header_cells = table.rows[0].cells
        for i, header_text in enumerate(element.headers):
            cell = header_cells[i]
            set_cell_shading(cell, PRIMARY_HEX)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header_text)
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            run.font.color.rgb = WHITE_COLOR

        # Filas de datos
        for row_idx, row_data in enumerate(element.rows):
            row = table.add_row()
            bg_color = LIGHT_GRAY_HEX if row_idx % 2 == 0 else 'FFFFFF'

            for i, cell_text in enumerate(row_data):
                if i < len(row.cells):
                    cell = row.cells[i]
                    set_cell_shading(cell, bg_color)
                    set_cell_margins(cell)
                    p = cell.paragraphs[0]
                    self._add_formatted_text(p, cell_text)
                    p.paragraph_format.space_after = Pt(0)

        # Espaciado después de tabla
        self.doc.add_paragraph()

    def _render_code_block(self, element: CodeBlockElement):
        """Renderiza bloque de código."""
        if self.ascii_as_image:
            self._render_code_as_image(element)
        else:
            self._render_code_as_text(element)

    def _render_code_as_text(self, element: CodeBlockElement):
        """Renderiza código como texto monoespaciado en tabla."""
        # Crear tabla de 1 celda como contenedor
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        set_cell_shading(cell, LIGHT_GRAY_HEX)
        set_table_borders(table, BORDER_GRAY_HEX)
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)

        # Limpiar párrafo existente
        cell.paragraphs[0].clear()

        for i, line in enumerate(element.lines):
            if i == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()

            # Preservar espacios usando non-breaking space si es necesario
            display_line = line if line else ' '
            run = p.add_run(display_line)
            run.font.name = FONT_NAME_MONO
            run.font.size = Pt(9)
            run.font.color.rgb = TEXT_COLOR
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0

        self.doc.add_paragraph()

    def _render_code_as_image(self, element: CodeBlockElement):
        """Renderiza código como imagen PNG."""
        try:
            from PIL import Image, ImageDraw, ImageFont
            import tempfile

            # Configuración
            font_size = 14
            padding = 20
            line_height = font_size + 4

            # Intentar cargar fuente monoespaciada
            try:
                font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf", font_size)
            except:
                font = ImageFont.load_default()

            # Calcular dimensiones
            max_width = max(len(line) for line in element.lines) if element.lines else 10
            width = max_width * (font_size * 0.6) + padding * 2
            height = len(element.lines) * line_height + padding * 2

            # Crear imagen
            img = Image.new('RGB', (int(width), int(height)), color='#F5F5F5')
            draw = ImageDraw.Draw(img)

            # Dibujar texto
            y = padding
            for line in element.lines:
                draw.text((padding, y), line, fill='#303030', font=font)
                y += line_height

            # Guardar temporalmente
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                img.save(tmp.name)

                # Insertar en documento
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(tmp.name, width=Cm(15))

                # Limpiar
                os.unlink(tmp.name)

            self.doc.add_paragraph()

        except ImportError:
            print("Pillow no instalado. Usando texto monoespaciado.")
            self._render_code_as_text(element)

    def _render_blockquote(self, element: BlockquoteElement):
        """Renderiza blockquote con borde naranja."""
        p = self.doc.add_paragraph()

        # Borde izquierdo naranja
        add_left_border(p, ACCENT_HEX, width='24')

        # Fondo claro
        set_paragraph_shading(p, BLOCKQUOTE_BG_HEX)

        # Indentación
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        # Texto
        self._add_formatted_text(p, element.text)

        self.doc.add_paragraph()

    def _render_image(self, element: ImageElement):
        """Renderiza imagen centrada con caption."""
        image_path = self.base_image_path / element.src

        if not image_path.exists():
            # Placeholder si no existe
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[Imagen no encontrada: {element.src}]")
            run.italic = True
            run.font.color.rgb = GRAY_COLOR
            return

        # Insertar imagen
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()

        # Ajustar ancho según el tipo de imagen
        width = Cm(15)
        if 'heatmap' in element.src.lower() and 'combined' not in element.src.lower():
            width = Cm(12)  # Heatmaps individuales más pequeños

        run.add_picture(str(image_path), width=width)

        # Caption
        if element.alt:
            p_caption = self.doc.add_paragraph()
            p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_caption.add_run(element.alt)
            run.italic = True
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            run.font.color.rgb = GRAY_COLOR

        self.doc.add_paragraph()

    def _render_list(self, element: ListElement):
        """Renderiza lista con viñetas o numerada."""
        style = 'List Number' if element.is_numbered else 'List Bullet'

        for item in element.items:
            p = self.doc.add_paragraph(style=style)
            self._add_formatted_text(p, item)
            p.paragraph_format.left_indent = Cm(1.27)

    def _render_horizontal_rule(self):
        """Renderiza línea horizontal decorativa."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run('─' * 30)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


# =============================================================================
# MAIN
# =============================================================================

def extract_metadata(content: str) -> dict:
    """Extrae metadatos del documento MD."""
    metadata = {
        'title': 'Sistema de Alerta Temprana v2',
        'subtitle': 'Predicción de Fracaso Académico mediante Patrones de Navegación en LMS',
        'date': '5 de enero de 2026',
        'version': '2.0'
    }

    # Buscar título
    match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
    if match:
        metadata['title'] = match.group(1).strip()

    # Buscar subtítulo (## después del título)
    match = re.search(r'^##\s+([^#].+)$', content, re.MULTILINE)
    if match:
        metadata['subtitle'] = match.group(1).strip()

    # Buscar fecha
    match = re.search(r'\*\*Fecha:\*\*\s*(.+)$', content, re.MULTILINE)
    if match:
        metadata['date'] = match.group(1).strip()

    # Buscar versión
    match = re.search(r'\*\*Versión:\*\*\s*(.+)$', content, re.MULTILINE)
    if match:
        metadata['version'] = match.group(1).strip()

    return metadata


def filter_markdown_toc(elements: List[Element]) -> List[Element]:
    """
    Filtra la sección de 'Tabla de Contenidos' del markdown original.
    Esta sección se genera automáticamente en el DOCX, así que la del MD es redundante.
    """
    filtered = []
    skip_until_next_heading = False

    for elem in elements:
        # Si encontramos "Tabla de Contenidos" como heading, empezamos a saltar
        if isinstance(elem, HeadingElement):
            if 'tabla de contenidos' in elem.text.lower():
                skip_until_next_heading = True
                continue
            else:
                # Encontramos otro heading, dejamos de saltar
                skip_until_next_heading = False

        # Si estamos en modo skip, ignorar elementos (listas de TOC)
        if skip_until_next_heading:
            continue

        filtered.append(elem)

    return filtered


def main():
    parser = argparse.ArgumentParser(description='Convertir MD a DOCX profesional')
    parser.add_argument('--input', type=Path, default=INPUT_FILE,
                        help='Archivo MD de entrada')
    parser.add_argument('--output', type=Path, default=OUTPUT_FILE,
                        help='Archivo DOCX de salida')
    parser.add_argument('--ascii-as-image', action='store_true',
                        help='Convertir bloques ASCII a imágenes PNG')
    args = parser.parse_args()

    print(f"Convirtiendo: {args.input}")
    print(f"Salida: {args.output}")
    print(f"ASCII como imagen: {args.ascii_as_image}")
    print()

    # Leer contenido
    with open(args.input, 'r', encoding='utf-8') as f:
        content = f.read()

    # Extraer metadatos
    metadata = extract_metadata(content)
    print(f"Título: {metadata['title']}")
    print(f"Subtítulo: {metadata['subtitle']}")
    print()

    # Parsear Markdown
    print("Parseando Markdown...")
    parser_md = MarkdownParser(content)
    elements = parser_md.parse()
    print(f"  Elementos parseados: {len(elements)}")

    # Filtrar la sección de TOC del markdown (se genera automáticamente en DOCX)
    elements = filter_markdown_toc(elements)
    print(f"  Elementos después de filtrar TOC: {len(elements)}")

    # Contar tipos
    type_counts = {}
    for elem in elements:
        t = elem.type.value
        type_counts[t] = type_counts.get(t, 0) + 1
    for t, c in sorted(type_counts.items()):
        print(f"    - {t}: {c}")
    print()

    # Crear documento
    print("Creando documento...")
    builder = DocumentBuilder(args.output)

    # Portada
    print("  - Agregando portada...")
    builder.add_title_page(
        title=metadata['title'],
        subtitle=metadata['subtitle'],
        date=metadata['date'],
        version=metadata['version']
    )

    # Tabla de contenidos
    print("  - Generando tabla de contenidos...")
    headings = [e for e in elements if isinstance(e, HeadingElement) and e.level <= 2]
    builder.add_table_of_contents(headings)

    # Headers y footers
    print("  - Configurando headers/footers...")
    builder.setup_headers_footers(metadata['title'])

    # Renderizar contenido
    print("  - Renderizando contenido...")
    renderer = ElementRenderer(
        builder.doc,
        REPORT_DIR,
        ascii_as_image=args.ascii_as_image
    )

    for i, element in enumerate(elements):
        renderer.render(element)
        if (i + 1) % 50 == 0:
            print(f"    Procesados {i + 1}/{len(elements)} elementos...")

    # Guardar
    print()
    print("Guardando documento...")
    builder.save()

    print()
    print(f"¡Documento generado exitosamente!")
    print(f"Archivo: {args.output}")

    # Verificar tamaño
    size_mb = args.output.stat().st_size / (1024 * 1024)
    print(f"Tamaño: {size_mb:.2f} MB")


if __name__ == '__main__':
    main()
